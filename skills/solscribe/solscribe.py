#!/usr/bin/env python3
"""
SolScribe — Book writing companion.
Manages chapters as .md files with YAML frontmatter.
Backs up locally. Git push is disabled for chapter content (private on-disk only).
"""

import json
import re
import uuid
import shutil
import subprocess
from pathlib import Path
from datetime import datetime

BASE_DIR = Path(__file__).parent
STATE_FILE = BASE_DIR / "book_state.json"
CHAPTERS_DIR = BASE_DIR / "chapters"
BACKUPS_DIR = BASE_DIR / "backups"
LOGS_DIR = BASE_DIR / "session_logs"


def slugify(title):
    """Make a URL-safe slug from a title."""
    slug = title.lower()
    slug = re.sub(r'[^\w\s-]', '', slug)
    slug = re.sub(r'[\s_]+', '-', slug)
    slug = slug.strip('-')
    return slug


def load_state():
    if not STATE_FILE.exists():
        return {"book_title": "My Memoir", "author": "Annmarie Lee", "chapters": []}
    with open(STATE_FILE) as f:
        return json.load(f)


def save_state(state):
    with open(STATE_FILE, "w") as f:
        json.dump(state, f, indent=2)


def chapter_file(chapter):
    """Return Path to chapter .md file."""
    order = str(chapter.get("order", 0)).zfill(3)
    slug = slugify(chapter["title"])[:40]
    return CHAPTERS_DIR / f"{order}-{slug}.md"


def read_chapter_content(chapter):
    """Read raw markdown content from chapter file."""
    path = chapter_file(chapter)
    if not path.exists():
        return ""
    text = path.read_text()
    # Strip frontmatter
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            return parts[2].strip()
    return text


def write_chapter_file(chapter, content):
    """Write chapter content with YAML frontmatter."""
    path = chapter_file(chapter)
    path.parent.mkdir(parents=True, exist_ok=True)
    frontmatter = f"""---
title: "{chapter['title']}"
order: {chapter.get('order', 0)}
status: {chapter.get('status', 'planned')}
created: {chapter.get('created', datetime.now().isoformat())}
updated: {datetime.now().isoformat()}
---

"""
    path.write_text(frontmatter + content.strip())


def count_words(chapter):
    text = read_chapter_content(chapter)
    words = text.split()
    return len(words)


def get_chapters():
    state = load_state()
    chapters = []
    for ch in state["chapters"]:
        ch = dict(ch)
        ch["word_count"] = count_words(ch)
        chapters.append(ch)
    return sorted(chapters, key=lambda c: c.get("order", 0))


def get_chapter(chapter_id):
    state = load_state()
    for ch in state["chapters"]:
        if ch["id"] == chapter_id:
            ch = dict(ch)
            ch["word_count"] = count_words(ch)
            ch["content"] = read_chapter_content(ch)
            return ch
    return None


def get_chapter_by_order(order):
    state = load_state()
    for ch in state["chapters"]:
        if ch.get("order") == order:
            return get_chapter(ch["id"])
    return None


def total_words():
    return sum(count_words(ch) for ch in load_state()["chapters"])


def create_chapter(title, content="", status="planned"):
    """
    Create a new chapter as a .md file.
    Chapters are NEVER overwritten. Each chapter is permanent.
    Use revise_chapter() to update a chapter's content.
    """
    state = load_state()
    chapter_id = str(uuid.uuid4())[:8]
    # Next order number — based on current highest order
    existing_orders = [c.get("order", 0) for c in state["chapters"]]
    n = max(existing_orders, default=0) + 1

    if not title:
        title = f"Chapter {n}"

    chapter = {
        "id": chapter_id,
        "title": title,
        "status": status,
        "created": datetime.now().isoformat(),
        "updated": datetime.now().isoformat(),
        "order": n,
    }

    state["chapters"].append(chapter)
    save_state(state)

    write_chapter_file(chapter, content.strip())
    backup(f"Create chapter: {title}")

    return chapter


def revise_chapter(chapter_id, content=None, title=None, status=None):
    """
    Revise an existing chapter with new content.
    The old version is preserved in backups.
    This replaces content while keeping the chapter's identity.
    """
    state = load_state()
    for ch in state["chapters"]:
        if ch["id"] == chapter_id:
            if title is not None:
                ch["title"] = title
            if status is not None:
                ch["status"] = status
            ch["updated"] = datetime.now().isoformat()
            break
    save_state(state)

    if content is not None:
        chapter = get_chapter(chapter_id)
        write_chapter_file(chapter, content)

    backup(f"Revise chapter: {chapter_id}")
    return get_chapter(chapter_id)


def update_chapter(chapter_id, title=None, content=None, status=None):
    """Update chapter title, content, or status."""
    state = load_state()
    for ch in state["chapters"]:
        if ch["id"] == chapter_id:
            if title is not None:
                ch["title"] = title
            if status is not None:
                ch["status"] = status
            ch["updated"] = datetime.now().isoformat()
            break
    save_state(state)

    if content is not None:
        chapter = get_chapter(chapter_id)
        write_chapter_file(chapter, content)

    backup(f"Update chapter: {title or chapter_id}")
    return get_chapter(chapter_id)


def delete_chapter(chapter_id):
    """Delete a chapter and its file."""
    state = load_state()
    chapter = next((c for c in state["chapters"] if c["id"] == chapter_id), None)
    title = chapter["title"] if chapter else chapter_id
    state["chapters"] = [c for c in state["chapters"] if c["id"] != chapter_id]
    save_state(state)

    if chapter:
        path = chapter_file(chapter)
        if path.exists():
            path.unlink()

    backup(f"Delete chapter: {title}")


def append_to_chapter(chapter_id, text):
    """Append text to existing chapter."""
    chapter = get_chapter(chapter_id)
    existing = chapter["content"] if chapter else ""
    new_content = existing + "\n\n" + text.strip()
    return update_chapter(chapter_id, content=new_content)


def parse_incoming(text, chapter_hint=None):
    """Parse incoming text. Returns (action, chapter_id, content)."""
    text = text.strip()
    if not text:
        return ("empty", None, None)

    state = load_state()

    new_match = re.match(r"^new chapter:?\s*(.*)$", text, re.IGNORECASE)
    if new_match:
        title = new_match.group(1).strip() or None
        content = re.sub(r"^new chapter:?\s*", "", text, flags=re.IGNORECASE).strip()
        return ("create", title, content)

    ch_match = re.match(r"^(?:chapter|ch)\s*(\d+)[:\s]", text, re.IGNORECASE)
    if ch_match:
        num = int(ch_match.group(1))
        for ch in state["chapters"]:
            if ch.get("order") == num:
                content = re.sub(r"^(?:chapter|ch)\s*\d+[:\s]*", "", text, flags=re.IGNORECASE).strip()
                return ("append", ch["id"], content)

    if chapter_hint:
        return ("append", chapter_hint, text)

    if len(state["chapters"]) == 0:
        return ("create", "Chapter 1", text)

    if len(state["chapters"]) == 1:
        return ("append", state["chapters"][0]["id"], text)

    return ("ask", None, None)


def set_book_meta(book_title=None, author=None):
    state = load_state()
    if book_title is not None:
        state["book_title"] = book_title
    if author is not None:
        state["author"] = author
    save_state(state)
    return state


def log_session(amre_message, sol_response):
    """Log a conversation exchange to a dated session file."""
    LOGS_DIR.mkdir(exist_ok=True)
    today = datetime.now().strftime("%Y-%m-%d")
    log_file = LOGS_DIR / f"{today}.md"
    entry = f"\n## {datetime.now().strftime('%H:%M:%S')}\n\n**Amre:** {amre_message}\n\n**Sol:** {sol_response}\n\n---\n"
    with open(log_file, "a") as f:
        f.write(entry)


def backup(change_desc=""):
    """Local timestamped backup. No cloud push — chapter content stays private."""
    BACKUPS_DIR.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    backup_dir = BACKUPS_DIR / ts
    backup_dir.mkdir(exist_ok=True)
    shutil.copy2(STATE_FILE, backup_dir / "book_state.json")
    for ch_file in CHAPTERS_DIR.glob("*.md"):
        shutil.copy2(ch_file, backup_dir / ch_file.name)
    backups = sorted(BACKUPS_DIR.iterdir())
    while len(backups) > 50:
        oldest = backups.pop(0)
        shutil.rmtree(oldest)


# ── DOCX Export ────────────────────────────────────────────────────────────────

def export_docx(output_path=None):
    """Export all chapters to DOCX."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed")

    state = load_state()
    doc = Document()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(state["book_title"])
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x24, 0x16)

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(state["author"])
    author_run.font.size = Pt(14)
    author_run.font.color.rgb = RGBColor(0x6B, 0x5D, 0x4D)

    doc.add_page_break()

    sorted_chapters = sorted(state["chapters"], key=lambda c: c.get("order", 0))
    for ch in sorted_chapters:
        full = get_chapter(ch["id"])
        heading = doc.add_heading(full["title"], level=1)
        heading.runs[0].font.color.rgb = RGBColor(0xB8, 0x5C, 0x38)
        content = full.get("content", "").strip()
        if content:
            doc.add_paragraph(content)
        if ch != sorted_chapters[-1]:
            doc.add_page_break()

    if output_path:
        doc.save(output_path)
    return doc
